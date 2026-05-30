import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc_path = r'c:\Users\akshi\Desktop\GIT_STAR\Spatial_Autocorrelation_Assignment.docx'
out_path = r'c:\Users\akshi\Desktop\GIT_STAR\dumped_doc.txt'

try:
    doc = docx.Document(doc_path)
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write("=== DOCUMENT DUMP ===\n\n")
        f.write(f"Number of paragraphs: {len(doc.paragraphs)}\n")
        f.write(f"Number of tables: {len(doc.tables)}\n")
        f.write(f"Number of sections: {len(doc.sections)}\n\n")
        
        # We walk through elements sequentially in document order
        # to preserve paragraph and table relative positions.
        body = doc.element.body
        for elem in body:
            # check if paragraph
            if elem.tag.endswith('p'):
                # find the paragraph object in doc.paragraphs
                p_id = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId')
                p_text = ""
                # Simple fallback: find by index or match
                # To be simple and robust:
                p = docx.text.paragraph.Paragraph(elem, doc)
                if p.text.strip():
                    f.write(f"[{p.style.name}] {p.text}\n\n")
            elif elem.tag.endswith('tbl'):
                t = docx.table.Table(elem, doc)
                f.write(f"--- TABLE (Rows: {len(t.rows)}, Cols: {len(t.columns)}) ---\n")
                for r in t.rows:
                    row_txt = [cell.text.strip().replace('\n', ' ') for cell in r.cells]
                    f.write(" | ".join(row_txt) + "\n")
                f.write("---------------------------------------\n\n")
    print(f"[+] Successfully dumped to {out_path}")
except Exception as e:
    print(f"[!] Error: {e}")
