import docx

def replace_text_in_doc(doc_path, output_path, replacements):
    doc = docx.Document(doc_path)
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        for search_text, replace_text in replacements.items():
            if search_text in p.text:
                # Naive replacement may lose formatting if the text is split across runs,
                # but if the text is simple, it's often fine. A more robust way:
                inline = p.runs
                for i in range(len(inline)):
                    if search_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(search_text, replace_text)
                # If it's still in paragraph text but not replaced (split across runs), replace at paragraph level and clear runs
                if search_text in p.text:
                    p.text = p.text.replace(search_text, replace_text)

    # Replace in tables
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for search_text, replace_text in replacements.items():
                        if search_text in p.text:
                            inline = p.runs
                            for i in range(len(inline)):
                                if search_text in inline[i].text:
                                    inline[i].text = inline[i].text.replace(search_text, replace_text)
                            if search_text in p.text:
                                p.text = p.text.replace(search_text, replace_text)
                                
    doc.save(output_path)
    print("Done")

replacements = {
    "N. Akshit Vinay": "TEST NAME",
    "25MSRSGIS001": "TEST PID",
    "M.Sc. GIS & Remote Sensing": "TEST COURSE"
}
replace_text_in_doc("demo_assign.docx", "test_python_docx.docx", replacements)
