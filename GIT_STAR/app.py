import os
import platform
import subprocess
from flask import Flask, render_template, request, send_file, jsonify
import docx
import shutil

app = Flask(__name__)

# Ensure absolute paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_FILE = os.path.join(BASE_DIR, "demo_assign.docx")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def replace_text_in_doc(doc_path, output_path, replacements):
    doc = docx.Document(doc_path)
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        for search_text, replace_text in replacements.items():
            if search_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if search_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(search_text, replace_text)
                if search_text in p.text:
                    p.text = p.text.replace(search_text, replace_text)

    # Replace in tables
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for search_text, replace_text in replacements.items():
                        if search_text in p.text:
                            inline = p.runs
                            for i in range(len(inline)):
                                if search_text in inline[i].text:
                                    inline[i].text = inline[i].text.replace(search_text, replace_text)
                            if search_text in p.text:
                                p.text = p.text.replace(search_text, replace_text)
                                
    doc.save(output_path)

def convert_to_pdf(docx_path, pdf_path):
    """
    Attempts to convert DOCX to PDF using LibreOffice first (cross-platform),
    and falls back to docx2pdf (Windows only).
    """
    try:
        cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path]
        if platform.system() == "Windows":
            cmd[0] = "soffice"
            
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode == 0 and os.path.exists(pdf_path):
            return True
    except Exception as e:
        print(f"LibreOffice conversion failed or not installed: {e}")
        
    # Fallback to docx2pdf
    try:
        from docx2pdf import convert
        import pythoncom
        pythoncom.CoInitialize()
        try:
            convert(docx_path, pdf_path)
        finally:
            pythoncom.CoUninitialize()
            
        if os.path.exists(pdf_path):
            return True
    except Exception as e:
        print(f"Fallback docx2pdf failed: {e}")
        
    return False

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate():
    data = request.json
    name = data.get("name", "N. Akshit Vinay")
    pid = data.get("pid", "25MSRSGIS001")
    course = data.get("course", "M.Sc. GIS & Remote Sensing")
    
    # Create safe filename prefix
    safe_name = "".join([c if c.isalnum() else "_" for c in name]).strip("_")
    safe_pid = "".join([c if c.isalnum() else "_" for c in pid]).strip("_")
    filename_prefix = f"{safe_name}_{safe_pid}"
        
    docx_filename = f"{filename_prefix}.docx"
    pdf_filename = f"{filename_prefix}.pdf"
    
    docx_path = os.path.join(OUTPUT_DIR, docx_filename)
    pdf_path = os.path.join(OUTPUT_DIR, pdf_filename)
    
    try:
        if not os.path.exists(TEMPLATE_FILE):
            return jsonify({"error": f"Template file not found at {TEMPLATE_FILE}"}), 500
            
        # Perform text replacement
        replacements = {
            "N. Akshit Vinay": name,
            "25MSRSGIS001": pid,
            "M.Sc. GIS & Remote Sensing": course
        }
        replace_text_in_doc(TEMPLATE_FILE, docx_path, replacements)
        
        # Convert to PDF
        print(f"Converting to PDF: {pdf_path}")
        success = convert_to_pdf(docx_path, pdf_path)
        
        if not success:
            return jsonify({"error": "Failed to convert document to PDF. Ensure LibreOffice or MS Word is installed."}), 500
        
        # Return the URL to access the PDF
        return jsonify({
            "success": True,
            "pdf_url": f"/download/{pdf_filename}"
        })
        
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/download/<filename>")
def download(filename):
    file_path = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=False)
    return "File not found", 404

if __name__ == "__main__":
    # Important to bind to 0.0.0.0 for Docker containers
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
