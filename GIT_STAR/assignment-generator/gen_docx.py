import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Georgia'
    return h

def create_doc():
    doc = Document()
    
    # --- COVER PAGE ---
    # Add logo
    logo_path = os.path.join(os.getcwd(), 'shuats_logo.png')
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_logo.add_run()
        r.add_picture(logo_path, width=Inches(1.5))
        
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_inst = p_inst.add_run('SAM HIGGINBOTTOM UNIVERSITY OF AGRICULTURE, TECHNOLOGY AND SCIENCES\n')
    r_inst.bold = True
    r_inst.font.size = Pt(14)
    r_sub = p_inst.add_run('SHUATS, Prayagraj - 211007, Uttar Pradesh, India\n\n')
    r_sub.font.size = Pt(11)

    p_tags = doc.add_paragraph()
    p_tags.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tags = p_tags.add_run('M.Sc. GIS & REMOTE SENSING (SEMESTER-II)\nCourse: MAS 744 — Geospatial Statistics- II\n\n')
    r_tags.bold = True
    r_tags.font.size = Pt(12)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('SPATIAL AUTOCORRELATION\n')
    r_title.bold = True
    r_title.font.size = Pt(28)
    r_subt = p_title.add_run('& Point Pattern Analysis\n\n\n')
    r_subt.italic = True
    r_subt.font.size = Pt(18)

    # Submitted details
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1 = table.cell(0, 0)
    c1p = cell1.paragraphs[0]
    c1p.add_run('SUBMITTED BY:\n').bold = True
    c1p.add_run('N. Akshit Vinay\nPID: 25MSRSGIS001\nProgram: M.Sc. GIS & Remote Sensing')
    
    cell2 = table.cell(0, 1)
    c2p = cell2.paragraphs[0]
    c2p.add_run('SUBMITTED TO:\n').bold = True
    c2p.add_run('Mis. Arpita Esther\nDept. of Mathematics & Statistics')

    doc.add_page_break()

    # --- TOC ---
    add_heading(doc, 'Table of Contents', 1)
    toc_text = """TOPIC 1: Spatial Autocorrelation, Moran's I & Getis-Ord Gi*
1.1 Spatial Autocorrelation — Definition & Types
1.2 Moran's I — The Core Measure
1.3 Getis-Ord Gi* — Local Hotspot Analysis

TOPIC 2: Point Pattern Analysis & Hotspot Detection
2.1 Point Pattern Analysis — Definition & Types
2.2 Method 1: Nearest Neighbour Analysis (NNA)
2.3 Method 2: Ripley's K Function
2.4 Hotspot Detection — Kernel Density Estimation
2.5 Complete Workflow in GIS"""
    p_toc = doc.add_paragraph(toc_text)
    
    doc.add_page_break()

    # --- TOPIC 1 ---
    add_heading(doc, 'TOPIC 1: SPATIAL AUTOCORRELATION', 1)
    
    add_heading(doc, '1.1 Spatial Autocorrelation', 2)
    p_def = doc.add_paragraph()
    p_def.add_run('Definition').bold = True
    doc.add_paragraph('Spatial Autocorrelation is a measure that tells us whether geographic features (like villages, rainfall stations, or land cover patches) that are close to each other are more similar or more different compared to features that are far apart. In simple words, it answers the question: Are nearby things similar to each other?')
    doc.add_paragraph('This concept is based on the First Law of Geography proposed by Waldo Tobler (1970), which states:')
    p_quote = doc.add_paragraph('"Everything is related to everything else, but near things are more related than distant things." — Waldo Tobler, 1970')

    p_type = doc.add_paragraph()
    p_type.add_run('Types of Spatial Autocorrelation').bold = True
    t1 = doc.add_table(rows=4, cols=2)
    t1.style = 'Table Grid'
    t1.cell(0,0).text = 'Type'
    t1.cell(0,1).text = 'Meaning & Explanation'
    t1.cell(1,0).text = 'Positive Autocorrelation'
    t1.cell(1,1).text = 'Similar values are found near each other. Example: High rainfall areas surrounded by other high rainfall areas. Most common pattern in nature.'
    t1.cell(2,0).text = 'No Autocorrelation (Random)'
    t1.cell(2,1).text = 'Values are scattered randomly with no spatial pattern. Example: Random distribution of soil sample values with no grouping.'
    t1.cell(3,0).text = 'Negative Autocorrelation'
    t1.cell(3,1).text = 'Dissimilar values are next to each other (checkerboard). Example: High crop yield areas surrounded by low yield areas due to soil variation.'

    p_why = doc.add_paragraph()
    p_why.add_run('\nWhy is Spatial Autocorrelation Important in GIS?').bold = True
    doc.add_paragraph('It tells us whether a spatial pattern is random or structured.', style='List Bullet')
    doc.add_paragraph('It helps choose the right statistical method for analysis.', style='List Bullet')
    doc.add_paragraph('Standard statistics assume data independence; ignoring spatial dependence gives wrong results.', style='List Bullet')
    doc.add_paragraph('It is the foundation for hotspot analysis, cluster detection, and spatial regression.', style='List Bullet')
    
    doc.add_page_break()

    # --- Moran's I ---
    add_heading(doc, "1.2 Moran's I — The Core Measure", 2)
    p_def2 = doc.add_paragraph()
    p_def2.add_run('Definition').bold = True
    doc.add_paragraph("Moran's I is the most widely used statistical measure to quantify spatial autocorrelation. It was developed by Patrick Alfred Pierce Moran in 1950. It gives a single number summarising whether values are clustered, dispersed, or random — a 'correlation coefficient for space.'")
    
    p_form = doc.add_paragraph("I = (N / W) × [ Σᵢ Σⱼ wᵢⱼ (xᵢ – x̄)(xⱼ – x̄) ] / [ Σᵢ (xᵢ – x̄)² ]")
    p_form.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_int = doc.add_paragraph()
    p_int.add_run('\nInterpretation of Moran\'s I').bold = True
    t2 = doc.add_table(rows=4, cols=3)
    t2.style = 'Table Grid'
    t2.cell(0,0).text = "Moran's I Value"
    t2.cell(0,1).text = "Interpretation"
    t2.cell(0,2).text = "GIS/RS Example"
    t2.cell(1,0).text = "Close to +1"
    t2.cell(1,1).text = "Strong Positive Autocorrelation (Clustering)"
    t2.cell(1,2).text = "NDVI values cluster — forests near forests, bare land near bare land"
    t2.cell(2,0).text = "Close to 0"
    t2.cell(2,1).text = "No Autocorrelation (Random Pattern)"
    t2.cell(2,2).text = "Randomly distributed soil pH with no spatial pattern"
    t2.cell(3,0).text = "Close to -1"
    t2.cell(3,1).text = "Strong Negative Autocorrelation (Dispersion)"
    t2.cell(3,2).text = "Checkerboard of high/low crop yield in mixed agriculture zone"

    doc.add_page_break()

    # --- Getis-Ord ---
    add_heading(doc, "1.3 Getis-Ord Gi* — Local Hotspot Analysis", 2)
    p_def3 = doc.add_paragraph()
    p_def3.add_run('Definition').bold = True
    doc.add_paragraph("The Getis-Ord Gi* statistic (pronounced 'Gee-I-Star') is a LOCAL spatial statistic that identifies specific locations which are statistically significant hotspots or coldspots. Unlike Moran's I which gives one result for the whole study area, Gi* gives a separate result (Z-score and p-value) for EVERY SINGLE location in the dataset. Developed by Getis and Ord in 1992.")

    p_int2 = doc.add_paragraph()
    p_int2.add_run('\nInterpretation of Gi* Results').bold = True
    t3 = doc.add_table(rows=4, cols=3)
    t3.style = 'Table Grid'
    t3.cell(0,0).text = "Gi* Z-score"
    t3.cell(0,1).text = "Type"
    t3.cell(0,2).text = "What It Means"
    t3.cell(1,0).text = "High Positive + p < 0.05"
    t3.cell(1,1).text = "HOTSPOT"
    t3.cell(1,2).text = "Location and neighbors have unusually HIGH values — significant high-value cluster"
    t3.cell(2,0).text = "High Negative + p < 0.05"
    t3.cell(2,1).text = "COLDSPOT"
    t3.cell(2,2).text = "Location and neighbors have unusually LOW values — significant low-value cluster"
    t3.cell(3,0).text = "Near 0 + p > 0.05"
    t3.cell(3,1).text = "Not Significant"
    t3.cell(3,2).text = "Pattern at this location is random. Cannot be called hotspot or coldspot."

    doc.add_page_break()

    # --- TOPIC 2 ---
    add_heading(doc, 'TOPIC 2: POINT PATTERN ANALYSIS', 1)
    
    add_heading(doc, "2.1 Point Pattern Analysis", 2)
    p_def4 = doc.add_paragraph()
    p_def4.add_run('Definition').bold = True
    doc.add_paragraph('Point Pattern Analysis (PPA) is a set of statistical methods used to examine the spatial distribution of discrete point locations. It studies whether a set of points (disease cases, tree locations, crime events, earthquake epicenters) are arranged in a RANDOM, CLUSTERED, or DISPERSED pattern across a geographic area.')

    p_int3 = doc.add_paragraph()
    p_int3.add_run('\nThree Types of Spatial Point Patterns').bold = True
    t4 = doc.add_table(rows=4, cols=3)
    t4.style = 'Table Grid'
    t4.cell(0,0).text = "Pattern"
    t4.cell(0,1).text = "Description"
    t4.cell(0,2).text = "Real-World GIS Example"
    t4.cell(1,0).text = "Random"
    t4.cell(1,1).text = "Points scattered with no specific order. Null hypothesis = CSR."
    t4.cell(1,2).text = "Random placement of lightning strike locations on flat terrain."
    t4.cell(2,0).text = "Clustered"
    t4.cell(2,1).text = "Points group in specific areas due to an underlying driver."
    t4.cell(2,2).text = "Disease cases near contaminated water; fire hotspots in dry forest regions."
    t4.cell(3,0).text = "Dispersed"
    t4.cell(3,1).text = "Points more evenly spaced than random — suggests competition or planning."
    t4.cell(3,2).text = "Mobile phone towers at regular distances; wells deliberately spaced in a field."

    doc.add_page_break()

    # --- NNA & Ripley ---
    add_heading(doc, "2.2 Method 1: Nearest Neighbour Analysis (NNA)", 2)
    doc.add_paragraph('NNA measures the average distance from each point to its nearest neighboring point and compares this to the expected distance under CSR (Complete Spatial Randomness).')

    add_heading(doc, "2.3 Method 2: Ripley's K Function", 2)
    doc.add_paragraph("Ripley's K function is an advanced multi-scale point pattern analysis method. Instead of only looking at nearest neighbor distance, it counts points within successively larger circle radii to detect clustering or dispersion at MULTIPLE spatial scales simultaneously. Developed by Brian Ripley in 1976.")

    doc.add_page_break()

    # --- KDE ---
    add_heading(doc, "2.4 Kernel Density Estimation (KDE)", 2)
    doc.add_paragraph("KDE creates a continuous smooth density surface (raster) from point data, showing where points are most densely concentrated. Each point spreads as a smooth 'kernel' (bell-shaped) and contributions are summed at every location to create a density map. Widely used in crime analysis, wildlife ecology, accident mapping, and disease surveillance.")

    add_heading(doc, "2.5 Complete Workflow in GIS", 2)
    doc.add_paragraph('Collect Point Data: GPS survey, digitizing, or imagery extraction.', style='List Bullet')
    doc.add_paragraph('Load into GIS: Import CSV/shapefile into ArcGIS Pro or QGIS.', style='List Bullet')
    doc.add_paragraph('Define Study Area: Set boundary polygon for analysis region.', style='List Bullet')
    doc.add_paragraph('Conduct NNA: Get NNR and check if pattern is random, clustered, or dispersed.', style='List Bullet')
    doc.add_paragraph("Apply Ripley's K: Multi-scale analysis — see at which distances clustering occurs.", style='List Bullet')
    doc.add_paragraph('Run KDE: Create visual density surface (raster) showing hotspot zones.', style='List Bullet')

    doc.add_page_break()

    # --- END ---
    p_end = doc.add_paragraph('\n\n\n— The End —\n')
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run('N. Akshit Vinay\n25MSRSGIS001 · M.Sc. GIS & Remote Sensing')
    r_end.bold = True

    output_path = os.path.join(os.getcwd(), 'Assignment_Akshit_Vinay.docx')
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == '__main__':
    create_doc()
